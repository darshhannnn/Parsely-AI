"""
Enhanced DOCX content extractor with structure preservation
"""

import re
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass

from ...core.models import DocumentContent, ExtractedContent
from ...core.interfaces import DocumentType
from ...core.exceptions import ContentExtractionError
from ...core.logging_utils import get_pipeline_logger
from ...core.utils import create_temp_file, cleanup_temp_file, calculate_content_hash, timing_decorator


@dataclass
class DOCXParagraph:
    """Information about a DOCX paragraph"""
    text: str
    style_name: str
    is_heading: bool
    heading_level: int
    is_list_item: bool
    has_formatting: bool
    metadata: Dict[str, Any] = None


@dataclass
class DOCXSection:
    """DOCX document section"""
    title: str
    content: str
    paragraphs: List[DOCXParagraph]
    level: int
    word_count: int


@dataclass
class DOCXStructure:
    """DOCX document structure information"""
    total_paragraphs: int
    total_words: int
    total_chars: int
    sections: List[DOCXSection]
    headings: List[Dict[str, Any]]
    tables_count: int
    images_count: int
    styles_used: List[str]


class EnhancedDOCXExtractor:
    """Enhanced DOCX content extractor with structure preservation"""
    
    def __init__(self):
        self.logger = get_pipeline_logger()
    
    @timing_decorator
    def extract_content(self, document: DocumentContent) -> ExtractedContent:
        """Extract content from DOCX with enhanced structure analysis"""
        
        self.logger.info("Starting enhanced DOCX content extraction")
        
        try:
            from docx import Document
            
            # Save content to temporary file
            temp_file = create_temp_file(suffix='.docx')
            
            try:
                with open(temp_file, 'wb') as f:
                    f.write(document.raw_content)
                
                # Load document
                doc = Document(temp_file)
                
                # Extract paragraphs with detailed analysis
                paragraphs_info = self._extract_paragraphs_with_analysis(doc)
                
                # Analyze document structure
                structure = self._analyze_document_structure(doc, paragraphs_info)
                
                # Extract sections based on headings
                sections_dict = self._extract_sections_from_structure(structure)
                
                # Combine all text content
                text_content = '\n\n'.join([p.text for p in paragraphs_info if p.text.strip()])
                
                # Enhanced metadata
                metadata = {
                    'paragraph_count': len(paragraphs_info),
                    'section_count': len(structure.sections),
                    'heading_count': len(structure.headings),
                    'total_words': structure.total_words,
                    'total_chars': structure.total_chars,
                    'tables_count': structure.tables_count,
                    'images_count': structure.images_count,
                    'styles_used': structure.styles_used,
                    'extraction_method': 'Enhanced python-docx',
                    'structure_analysis': {
                        'sections': [
                            {
                                'title': s.title,
                                'level': s.level,
                                'word_count': s.word_count,
                                'paragraph_count': len(s.paragraphs)
                            }
                            for s in structure.sections
                        ],
                        'headings': structure.headings
                    }
                }
                
                # Add document properties if available
                if hasattr(doc, 'core_properties'):
                    props = doc.core_properties
                    metadata['document_properties'] = {
                        'title': props.title,
                        'author': props.author,
                        'subject': props.subject,
                        'keywords': props.keywords,
                        'category': props.category,
                        'comments': props.comments,
                        'created': props.created.isoformat() if props.created else None,
                        'modified': props.modified.isoformat() if props.modified else None,
                        'last_modified_by': props.last_modified_by,
                        'revision': props.revision,
                        'version': props.version
                    }
                
                return ExtractedContent(
                    document_id=calculate_content_hash(document.raw_content),
                    document_type=DocumentType.DOCX.value,
                    text_content=text_content,
                    pages=None,  # DOCX doesn't have fixed pages
                    sections=sections_dict,
                    metadata=metadata
                )
            
            finally:
                cleanup_temp_file(temp_file)
        
        except ImportError:
            raise ContentExtractionError("python-docx library not available for DOCX processing", DocumentType.DOCX.value)
        except Exception as e:
            raise ContentExtractionError(f"Enhanced DOCX extraction failed: {e}", DocumentType.DOCX.value)
    
    def _extract_paragraphs_with_analysis(self, doc) -> List[DOCXParagraph]:
        """Extract paragraphs with detailed analysis"""
        paragraphs_info = []
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            # Analyze paragraph style and formatting
            style_name = para.style.name if para.style else "Normal"
            is_heading = style_name.startswith('Heading')
            heading_level = self._extract_heading_level(style_name) if is_heading else 0
            is_list_item = self._is_list_item(para)
            has_formatting = self._has_formatting(para)
            
            # Extract additional metadata
            metadata = {
                'alignment': str(para.alignment) if para.alignment else None,
                'runs_count': len(para.runs),
                'has_hyperlinks': self._has_hyperlinks(para),
                'font_info': self._extract_font_info(para)
            }
            
            paragraph_info = DOCXParagraph(
                text=para.text.strip(),
                style_name=style_name,
                is_heading=is_heading,
                heading_level=heading_level,
                is_list_item=is_list_item,
                has_formatting=has_formatting,
                metadata=metadata
            )
            
            paragraphs_info.append(paragraph_info)
        
        return paragraphs_info
    
    def _analyze_document_structure(self, doc, paragraphs_info: List[DOCXParagraph]) -> DOCXStructure:
        """Analyze overall document structure"""
        
        total_paragraphs = len(paragraphs_info)
        total_words = sum(len(p.text.split()) for p in paragraphs_info)
        total_chars = sum(len(p.text) for p in paragraphs_info)
        
        # Count tables and images
        tables_count = len(doc.tables)
        images_count = self._count_images(doc)
        
        # Extract unique styles
        styles_used = list(set(p.style_name for p in paragraphs_info))
        
        # Build sections based on headings
        sections = self._build_sections(paragraphs_info)
        
        # Extract headings information
        headings = [
            {
                'text': p.text,
                'level': p.heading_level,
                'style': p.style_name,
                'position': i
            }
            for i, p in enumerate(paragraphs_info)
            if p.is_heading
        ]
        
        return DOCXStructure(
            total_paragraphs=total_paragraphs,
            total_words=total_words,
            total_chars=total_chars,
            sections=sections,
            headings=headings,
            tables_count=tables_count,
            images_count=images_count,
            styles_used=styles_used
        )
    
    def _build_sections(self, paragraphs_info: List[DOCXParagraph]) -> List[DOCXSection]:
        """Build sections based on heading structure"""
        sections = []
        current_section = None
        section_paragraphs = []
        
        for para in paragraphs_info:
            if para.is_heading:
                # Save previous section
                if current_section and section_paragraphs:
                    content = '\n\n'.join([p.text for p in section_paragraphs if not p.is_heading])
                    sections.append(DOCXSection(
                        title=current_section,
                        content=content,
                        paragraphs=section_paragraphs.copy(),
                        level=section_paragraphs[0].heading_level if section_paragraphs else 1,
                        word_count=len(content.split()) if content else 0
                    ))
                
                # Start new section
                current_section = para.text
                section_paragraphs = [para]
            else:
                if current_section:
                    section_paragraphs.append(para)
                else:
                    # Handle content before first heading
                    if not sections:
                        if not current_section:
                            current_section = "Introduction"
                            section_paragraphs = []
                        section_paragraphs.append(para)
        
        # Add final section
        if current_section and section_paragraphs:
            content = '\n\n'.join([p.text for p in section_paragraphs if not p.is_heading])
            sections.append(DOCXSection(
                title=current_section,
                content=content,
                paragraphs=section_paragraphs,
                level=section_paragraphs[0].heading_level if any(p.is_heading for p in section_paragraphs) else 1,
                word_count=len(content.split()) if content else 0
            ))
        
        return sections
    
    def _extract_sections_from_structure(self, structure: DOCXStructure) -> Dict[str, str]:
        """Extract sections dictionary from structure"""
        sections_dict = {}
        
        for section in structure.sections:
            sections_dict[section.title] = section.content
        
        return sections_dict
    
    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name"""
        if not style_name.startswith('Heading'):
            return 0
        
        # Extract number from "Heading 1", "Heading 2", etc.
        match = re.search(r'Heading\s*(\d+)', style_name)
        if match:
            return int(match.group(1))
        
        return 1  # Default to level 1
    
    def _is_list_item(self, para) -> bool:
        """Check if paragraph is a list item"""
        try:
            # Check if paragraph has numbering
            if para._element.pPr is not None:
                numPr = para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                if numPr is not None:
                    return True
            
            # Check for bullet-like characters at the start
            text = para.text.strip()
            if text and text[0] in '•·▪▫◦‣⁃':
                return True
            
            # Check for numbered list patterns
            if re.match(r'^\d+\.?\s+', text):
                return True
            
            return False
        except:
            return False
    
    def _has_formatting(self, para) -> bool:
        """Check if paragraph has special formatting"""
        try:
            for run in para.runs:
                if run.bold or run.italic or run.underline:
                    return True
                if run.font.color and run.font.color.rgb:
                    return True
                if run.font.highlight_color:
                    return True
            return False
        except:
            return False
    
    def _has_hyperlinks(self, para) -> bool:
        """Check if paragraph contains hyperlinks"""
        try:
            # Look for hyperlink elements in the paragraph
            hyperlinks = para._element.xpath('.//w:hyperlink', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            return len(hyperlinks) > 0
        except:
            return False
    
    def _extract_font_info(self, para) -> Dict[str, Any]:
        """Extract font information from paragraph"""
        font_info = {
            'fonts_used': [],
            'sizes_used': [],
            'has_bold': False,
            'has_italic': False,
            'has_underline': False
        }
        
        try:
            for run in para.runs:
                if run.font.name:
                    font_info['fonts_used'].append(run.font.name)
                if run.font.size:
                    font_info['sizes_used'].append(str(run.font.size))
                if run.bold:
                    font_info['has_bold'] = True
                if run.italic:
                    font_info['has_italic'] = True
                if run.underline:
                    font_info['has_underline'] = True
            
            # Remove duplicates
            font_info['fonts_used'] = list(set(font_info['fonts_used']))
            font_info['sizes_used'] = list(set(font_info['sizes_used']))
            
        except:
            pass
        
        return font_info
    
    def _count_images(self, doc) -> int:
        """Count images in the document"""
        try:
            image_count = 0
            
            # Count inline shapes (images)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run._element.xpath('.//pic:pic', namespaces={'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}):
                        image_count += 1
            
            # Count images in headers/footers if accessible
            for section in doc.sections:
                try:
                    header = section.header
                    footer = section.footer
                    
                    for container in [header, footer]:
                        for paragraph in container.paragraphs:
                            for run in paragraph.runs:
                                if run._element.xpath('.//pic:pic', namespaces={'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}):
                                    image_count += 1
                except:
                    continue
            
            return image_count
        except:
            return 0